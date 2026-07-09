import streamlit as st
import base64
import io
import os
import re
import json
from datetime import datetime, date, time as dt_time, timedelta
from pathlib import Path
import data_manager as dm
from PIL import Image
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
def create_docx(markdown_text: str) -> io.BytesIO:
    """Convert simple markdown text to a DOCX file."""
    doc = Document()
    
    # Configure default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for line in markdown_text.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
            
        if line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            # Basic bolding `**text**` replacement for simple cases
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def image_to_base64(uploaded_file) -> str:
    """Convert Streamlit UploadedFile to base64 JPEG string."""
    img = Image.open(uploaded_file)
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=90)
    return base64.b64encode(buf.getvalue()).decode("utf-8")


def resize_preview(uploaded_file, max_width: int = 600):
    """Resize image for preview representation."""
    img = Image.open(uploaded_file)
    w, h = img.size
    if w > max_width:
        img = img.resize((max_width, int(h * max_width / w)), Image.Resampling.LANCZOS)
    return img


# ─────────────────────────────────────────────────────────────────────────────
# AI BACKEND
# ─────────────────────────────────────────────────────────────────────────────

def _get_openai_client(api_key: str):
    from openai import OpenAI
    return OpenAI(api_key=api_key)


def call_text_ai(prompt: str, api_key: str, provider: str) -> str:
    """Run text-based model request."""
    if st.session_state.get("provider", "OpenAI (GPT-4o)") == "OpenAI (GPT-4o)":
        client = _get_openai_client(api_key)
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
        )
        return response.choices[0].message.content.strip()
    else:
        import google.generativeai as genai
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-1.5-pro")
        response = model.generate_content(prompt)
        return response.text.strip()


def call_vision_ai(prompt: str, image_b64: str, api_key: str, provider: str) -> str:
    """Run image + text request."""
    if st.session_state.get("provider", "OpenAI (GPT-4o)") == "OpenAI (GPT-4o)":
        client = _get_openai_client(api_key)
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{image_b64}",
                                "detail": "high",
                            },
                        },
                        {"type": "text", "text": prompt},
                    ],
                }
            ],
            temperature=0.6,
            max_tokens=4096,
        )
        return response.choices[0].message.content.strip()
    else:
        import google.generativeai as genai
        from google.generativeai.types import HarmCategory, HarmBlockThreshold
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-1.5-pro")
        image_bytes = base64.b64decode(image_b64)
        image_part = {"mime_type": "image/jpeg", "data": image_bytes}
        response = model.generate_content(
            [prompt, image_part],
            safety_settings={
                HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
                HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
            },
        )
        return response.text.strip()


# ─────────────────────────────────────────────────────────────────────────────
# SIDEBAR
# ─────────────────────────────────────────────────────────────────────────────
def render_sidebar():
    with st.sidebar:
        st.markdown(
            """
            <div style='text-align:center; padding:18px 0 10px 0;'>
                <span style='font-size:3rem;'>📚</span>
                <h2 style='margin:8px 0 2px 0; font-family: "Montserrat", sans-serif; font-size:1.35rem; font-weight:800;'>
                    Teacher Panel
                </h2>
                <p style='color:#94a3b8; font-size:0.8rem; margin:0;'>AI English Teaching Toolkit</p>
            </div>
            """,
            unsafe_allow_html=True,
        )
        st.divider()
    
        provider = st.selectbox(
            "🤖 AI Provider",
            ["OpenAI (GPT-4o)", "Google Gemini (1.5 Pro)"],
            help="Select model endpoint backend.",
        )
        from dotenv import load_dotenv
        load_dotenv()
        
        env_api_key = os.environ.get("OPENAI_API_KEY", "")
        if env_api_key:
            st.session_state.api_key = env_api_key
            st.success("🤖 API Key: Подключен (из .env)")
        else:
            api_key = st.text_input(
                "🔑 API Key",
                type="password",
                placeholder="sk-... or AIza...",
                help="Your secret key for OpenAI or Gemini APIs.",
            ) if "api_key" not in st.session_state else st.text_input(
                "🔑 API Key",
                value=st.session_state.get("api_key", ""),
                type="password",
                placeholder="sk-... or AIza...",
                help="Your secret key for OpenAI or Gemini APIs.",
            )
            if api_key:
                st.session_state.api_key = api_key
    
        student_level = st.selectbox(
            "🎓 Student Level",
            ["A1 – Beginner", "A2 – Elementary", "B1 – Intermediate",
             "B2 – Upper-Intermediate", "C1 – Advanced", "C2 – Proficiency"],
            index=3,
            help="Instructs the model to generate content appropriate for this level.",
        )
        st.session_state.level_code = student_level.split("–")[0].strip()
        st.session_state.provider = provider
        
        # ─── STUDENT MANAGEMENT ──────────────────────────────────────────────
        st.divider()
        st.markdown("#### 👥 Students")
        
        existing_students = dm.get_students()
        
        # Show payment warnings at the top
        pay_summaries = dm.get_all_payment_summaries()
        urgent = [p for p in pay_summaries if p["balance"] <= 0 and p["conducted"] > 0]
        if urgent:
            for u in urgent:
                st.error(f"💸 **{u['name']}** — пора платить! (баланс: {u['balance']})", icon="🔴")
    
        # Add student form
        with st.expander("➕ Add student", expanded=False):
            new_name = st.text_input("Name", placeholder="e.g. Настя", key="new_student_name")
            new_level = st.selectbox("Level", ["A1", "A2", "B1", "B2", "C1", "C2"], index=3, key="new_student_level")
            tz_options = ["Asia/Almaty", "Europe/Moscow", "Europe/London", "Europe/Berlin", "Asia/Tbilisi", "Asia/Yekaterinburg", "Asia/Dubai"]
            new_tz = st.selectbox("Student Timezone", tz_options, index=0, key="new_student_tz")
            if st.button("✅ Add", key="add_student_btn", use_container_width=True):
                if new_name.strip():
                    dm.add_student(new_name, new_level, new_tz)
                    st.success(f"Added: {new_name}")
                    st.rerun()
                else:
                    st.warning("Enter a name first")
    
        # List students
        if existing_students:
            pay_map = {p["student_id"]: p["balance"] for p in pay_summaries}
            for s in existing_students:
                balance = pay_map.get(s["id"], 0)
                bal_icon = "🟢" if balance > 2 else ("🟡" if balance > 0 else "🔴")
                tg_user = s.get("telegram_username", "")
                tg_badge = f" `@{tg_user}`" if tg_user else ""
                col_name, col_tg, col_del = st.columns([3, 1, 1])
                with col_name:
                    st.markdown(f"{bal_icon} **{s['name']}** ({s['level']}){tg_badge}")
                with col_tg:
                    tg_input = st.text_input(
                        "TG", value=tg_user, key=f"tg_{s['id']}",
                        placeholder="@user", label_visibility="collapsed",
                    )
                    if tg_input != tg_user:
                        dm.link_telegram(s["id"], tg_input)
                        st.rerun()
                with col_del:
                    if st.button("🗑️", key=f"del_student_{s['id']}"):
                        dm.delete_student(s["id"])
                        st.rerun()
        else:
            st.caption("No students yet. Add one above.")
    
        # ─── NOTION INTEGRATION ──────────────────────────────────────────────
        st.divider()
        st.markdown("#### 📓 Notion Sync")
        
        env_notion_token = os.environ.get("NOTION_TOKEN", "")
        env_notion_db = os.environ.get("NOTION_DB_ID", "")
        
        if env_notion_token and env_notion_db:
            st.session_state.notion_token = env_notion_token
            st.session_state.notion_db_id = env_notion_db
            st.success("📓 Notion: Подключено (из .env)")
        else:
            notion_token = st.text_input(
                "Notion Token",
                type="password",
                placeholder="secret_...",
                key="notion_token",
                help="Create an integration at notion.so/my-integrations",
            )
            notion_db_id = st.text_input(
                "Database ID",
                placeholder="e.g. abc123def456...",
                key="notion_db_id",
                help="Copy the 32-character ID from your Notion database URL",
            )
            
            if notion_token and notion_db_id:
                if st.button("🔗 Test Connection", key="test_notion_btn", use_container_width=True):
                    ok, msg = notion_sync.test_notion_connection(notion_token, notion_db_id)
                    if ok:
                        st.success(msg)
                    else:
                        st.error(msg)
    
        # ─── TELEGRAM USERBOT ────────────────────────────────────────────────
        st.divider()
        st.markdown("#### 🤖 Telegram Bot")
        
        st.info("Конфигурация Telegram-бота теперь полностью перенесена в файл `.env` для безопасности.")
        
        st.caption(
            "⚡ Запуск бота: `python telegram_bot.py`\n"
            "Бот работает как твой приватный ассистент и отвечает только на твой TEACHER_TELEGRAM_ID."
        )
    
        st.divider()
        st.markdown(
            "<div style='font-size:0.75rem; color:#475569; text-align:center;'>English Teacher Dashboard v2.0</div>",
            unsafe_allow_html=True,
        )
    
    
# ─────────────────────────────────────────────────────────────────────────────
# API KEY GUARD
# ─────────────────────────────────────────────────────────────────────────────
def require_api_key() -> bool:
    if not st.session_state.get("api_key"):
        st.warning("⚠️ Please provide an API key in the sidebar to start utilizing AI features.", icon="🔑")
        return False
    return True
    
    
    # ─────────────────────────────────────────────────────────────────────────────
    # MAIN HEADER
    # ─────────────────────────────────────────────────────────────────────────────
    st.markdown(
        """
        <div style='text-align:center; padding: 18px 0 10px 0;'>
            <h1 style='font-size:2.8rem; font-weight:800; margin:0;
                       background: linear-gradient(135deg, #a78bfa 0%, #818cf8 50%, #f472b6 100%);
                       -webkit-background-clip:text; -webkit-text-fill-color:transparent;
                       background-clip:text;'>
                📚 English Teacher Dashboard
            </h1>
            <p style='color:#94a3b8; font-size:1.05rem; margin:8px 0 0 0;'>
                Advanced toolkit to correct grammar, build quizzes, extract vocabulary, and plan lessons
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )
    
    # ─────────────────────────────────────────────────────────────────────────────
    # TABS SETUP
    # ─────────────────────────────────────────────────────────────────────────────

